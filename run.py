from docx import Document
from doctopdf import convert_to

def main():
    template_file_path = 'idcard.docx'
    output_file_path = 'result.docx'

    variables = {
        "${NAME}": "Hossein",
        "${TITLE}": "Software Engineer",
        "${BIRTHDAY}": "02 Jan, 1991",
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    convert_to('./', output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()